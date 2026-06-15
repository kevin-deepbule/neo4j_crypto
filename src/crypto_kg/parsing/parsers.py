"""Parsers for Word, PDF and plain text evaluation reports."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from crypto_kg.models import ParsedDocument, Section, Table, normalize_cell

SECTION_RE = re.compile(
    r"^\s*(?P<no>(?:第[一二三四五六七八九十百]+章)|(?:\d+(?:\.\d+){0,5}))"
    r"[\s、．.]*"
    r"(?P<title>[^。；;:：\n]{0,80})\s*$"
)


def parse_document(path: str | Path) -> ParsedDocument:
    """Parse a Word/PDF/text document into sections and tables."""

    doc_path = Path(path)
    suffix = doc_path.suffix.lower()
    if suffix == ".docx":
        return parse_docx(doc_path)
    if suffix == ".pdf":
        return parse_pdf(doc_path)
    if suffix in {".txt", ".md", ".text"}:
        return parse_text(doc_path)
    raise ValueError(f"Unsupported document type: {doc_path.suffix}")


def parse_text(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8-sig")
    sections = build_sections(text.splitlines())
    document = ParsedDocument(path=path, sections=sections, metadata={"parser": "text"})
    document.tables = [table for section in sections for table in section.tables]
    return document


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise RuntimeError("python-docx is required for .docx parsing") from exc

    document = Document(str(path))
    sections: list[Section] = []
    current: Section | None = None

    for block in iter_docx_blocks(document, Paragraph, DocxTable):
        if isinstance(block, Paragraph):
            text = normalize_cell(block.text)
            if not text:
                continue
            heading = parse_heading(text, style_name=getattr(block.style, "name", ""))
            if heading:
                section_no, title, level = heading
                current = Section(section_no=section_no, title=title, level=level)
                sections.append(current)
            else:
                current = ensure_section(sections, current)
                current.content = append_text(current.content, text)
        elif isinstance(block, DocxTable):
            current = ensure_section(sections, current)
            rows = [[normalize_cell(cell.text) for cell in row.cells] for row in block.rows]
            table = Table(rows=dedupe_docx_rows(rows), section_no=current.section_no, section_title=current.title)
            current.tables.append(table)

    parsed = ParsedDocument(path=path, sections=sections, metadata={"parser": "python-docx"})
    parsed.tables = [table for section in sections for table in section.tables]
    return parsed


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        return parse_pdf_with_pdfplumber(path)
    except ImportError:
        return parse_pdf_with_pymupdf(path)


def parse_pdf_with_pdfplumber(path: Path) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise exc

    lines: list[str] = []
    extracted_tables: list[tuple[int, list[list[str]]]] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            lines.extend(text.splitlines())
            for table in page.extract_tables() or []:
                rows = [[normalize_cell(cell) for cell in row] for row in table]
                if rows:
                    extracted_tables.append((page_index, rows))

    sections = build_sections(lines)
    attach_pdf_tables(sections, extracted_tables)
    parsed = ParsedDocument(path=path, sections=sections, metadata={"parser": "pdfplumber"})
    parsed.tables = [table for section in sections for table in section.tables]
    return parsed


def parse_pdf_with_pymupdf(path: Path) -> ParsedDocument:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("pdfplumber or PyMuPDF(fitz) is required for .pdf parsing") from exc

    lines: list[str] = []
    with fitz.open(str(path)) as pdf:
        for page in pdf:
            lines.extend((page.get_text("text") or "").splitlines())
    sections = build_sections(lines)
    parsed = ParsedDocument(path=path, sections=sections, metadata={"parser": "pymupdf"})
    parsed.tables = []
    return parsed


def build_sections(lines: Iterable[str]) -> list[Section]:
    sections: list[Section] = []
    current: Section | None = None
    pending_table: list[list[str]] = []

    for raw_line in lines:
        line = normalize_cell(raw_line)
        if not line:
            flush_text_table(current, pending_table)
            pending_table = []
            continue

        heading = parse_heading(line)
        if heading:
            flush_text_table(current, pending_table)
            pending_table = []
            section_no, title, level = heading
            current = Section(section_no=section_no, title=title, level=level)
            sections.append(current)
            continue

        current = ensure_section(sections, current)
        if pending_table and is_markdown_separator(line):
            continue
        table_row = parse_text_table_row(line)
        if table_row:
            pending_table.append(table_row)
        else:
            flush_text_table(current, pending_table)
            pending_table = []
            current.content = append_text(current.content, line)

    flush_text_table(current, pending_table)
    return sections


def parse_heading(text: str, style_name: str = "") -> tuple[str, str, int] | None:
    style_level = heading_level_from_style(style_name)
    match = SECTION_RE.match(text)
    if not match:
        return None
    section_no = match.group("no").replace("．", ".")
    title = normalize_cell(match.group("title"))
    level = style_level or section_no.count(".") + 1
    return section_no, title, level


def heading_level_from_style(style_name: str) -> int | None:
    match = re.search(r"Heading\s+(\d+)|标题\s*(\d+)", style_name or "", flags=re.I)
    if not match:
        return None
    return int(match.group(1) or match.group(2))


def parse_text_table_row(line: str) -> list[str] | None:
    if "|" in line:
        cells = [normalize_cell(cell) for cell in line.strip("|").split("|")]
    elif "\t" in line:
        cells = [normalize_cell(cell) for cell in line.split("\t")]
    else:
        return None
    if len(cells) < 2 or set(cells) <= {"", "-", "---"}:
        return None
    return cells


def is_markdown_separator(line: str) -> bool:
    if "|" not in line:
        return False
    cells = [cell.strip() for cell in line.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def flush_text_table(current: Section | None, rows: list[list[str]]) -> None:
    if current is None or len(rows) < 2:
        return
    current.tables.append(Table(rows=rows.copy(), section_no=current.section_no, section_title=current.title))


def ensure_section(sections: list[Section], current: Section | None) -> Section:
    if current is not None:
        return current
    current = Section(section_no="0", title="未分章节", level=0)
    sections.append(current)
    return current


def append_text(existing: str, text: str) -> str:
    return f"{existing}\n{text}".strip() if existing else text


def iter_docx_blocks(document, paragraph_type, table_type):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield paragraph_type(child, document)
        elif child.tag.endswith("}tbl"):
            yield table_type(child, document)


def dedupe_docx_rows(rows: list[list[str]]) -> list[list[str]]:
    cleaned: list[list[str]] = []
    for row in rows:
        if cleaned and row == cleaned[-1]:
            continue
        cleaned.append(row)
    return cleaned


def attach_pdf_tables(sections: list[Section], tables: list[tuple[int, list[list[str]]]]) -> None:
    current = ensure_section(sections, sections[-1] if sections else None)
    for page, rows in tables:
        current.tables.append(Table(rows=rows, section_no=current.section_no, section_title=current.title, page=page))
